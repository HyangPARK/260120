import json
from datetime import date

import streamlit as st

from io import BytesIO
from docx import Document
from docx.shared import Pt


# 이미지 아이콘을 쓰고 싶다면 PIL이 필요합니다(없어도 앱은 동작)
try:
    from PIL import Image
except Exception:
    Image = None


def safe_load_image(path: str):
    """이미지 파일이 없거나 PIL이 없으면 None 반환."""
    if Image is None:
        return None
    try:
        return Image.open(path)
    except Exception:
        return None


# -----------------------------
# Page Config (아이콘: 이미지 -> 없으면 기본 이모지)
# -----------------------------
APP_ICON_PATH = "assets/app_icon.png"  # 있으면 자동 적용
app_icon = safe_load_image(APP_ICON_PATH)

st.set_page_config(
    page_title="자기소개서 웹앱",
    page_icon=app_icon if app_icon else "🧾",  # 이미지 없으면 🧾
    layout="wide",
)

# -----------------------------
# Header (항상 보이도록: st.*만 사용)
# -----------------------------
st.title("자기소개서 웹앱")
st.caption("입력 → 미리보기 → JSON 내보내기")

# 아이콘 이미지가 있으면 상단에 표시(없으면 안내만)
if app_icon:
    st.image(app_icon, width=72)
else:
    st.info("아이콘 이미지를 쓰려면 프로젝트에 assets/app_icon.png 파일을 추가하세요. (없어도 앱은 정상 실행됩니다)")

st.divider()

# -----------------------------
# Sidebar
# -----------------------------
with st.sidebar:
    st.header("⚙️ 옵션")
    preview_mode = st.selectbox("미리보기 스타일", ["카드형", "문서형"], index=0)
    show_json = st.checkbox("입력값(JSON) 보기", value=False)


# -----------------------------
# Layout
# -----------------------------
left, right = st.columns([1.1, 1.2], gap="large")

# -----------------------------
# Inputs
# -----------------------------
with left:
    st.subheader("1) 기본 정보")
    c1, c2 = st.columns(2)
    with c1:
        name = st.text_input("이름", value="홍길동")
        role = st.text_input("직함/역할", value="교육공학 연구자")
        email = st.text_input("이메일", value="your@email.com")
    with c2:
        location = st.text_input("지역", value="Seoul, KR")
        phone = st.text_input("연락처(선택)", value="")
        website = st.text_input("웹사이트(선택)", value="")

    st.divider()

    st.subheader("2) 자기소개서 본문")
    headline = st.text_input("한 줄 핵심 메시지", value="데이터 기반으로 학습경험을 설계하고 개선합니다.")
    intro = st.text_area(
        "자기소개(요약 단락)",
        value=(
            "저는 대학 원격수업과 디지털 기반 교육을 설계·운영하며, "
            "학습자 경험과 학습분석을 연결하는 연구를 수행해왔습니다. "
            "생성형 AI를 교육에 적용할 때 신뢰와 투명성, 학습성과를 함께 고려합니다."
        ),
        height=140,
    )

    st.divider()

    st.subheader("3) 강점/역량")
    strengths = st.text_area(
        "강점(불릿 형태 추천)",
        value="- 수업설계 및 e-Learning 콘텐츠 기획\n- 학습자 피드백 기반 개선\n- 텍스트마이닝/학습데이터 분석",
        height=120,
    )

    skills_raw = st.text_input("스킬 태그(콤마로 구분)", value="교육공학, K-MOOC, 학습분석, 텍스트마이닝, 생성형AI")
    skills = [s.strip() for s in skills_raw.split(",") if s.strip()]

    st.divider()

    st.subheader("4) 경험(경력/프로젝트) 최대 4개")
    experiences = []
    for i in range(1, 5):
        with st.expander(f"경험 {i} 입력", expanded=(i == 1)):
            title = st.text_input(f"제목 {i}", value=("AI 기반 원격수업 고도화" if i == 1 else ""), key=f"t{i}")
            org = st.text_input(f"기관/팀 {i}", value=("스마트융합교육센터" if i == 1 else ""), key=f"o{i}")
            d_from = st.date_input(f"시작일 {i}", value=date(2025, 1, 1) if i == 1 else date.today(), key=f"f{i}")
            d_to = st.date_input(f"종료일 {i}", value=date.today(), key=f"to{i}")
            detail = st.text_area(
                f"설명 {i}",
                value=("실습 중심 콘텐츠 설계, 운영 데이터 분석, 개선 루프 구축" if i == 1 else ""),
                height=90,
                key=f"d{i}",
            )

        if title.strip():
            experiences.append(
                {
                    "title": title.strip(),
                    "org": org.strip(),
                    "from": str(d_from),
                    "to": str(d_to),
                    "detail": detail.strip(),
                }
            )

    st.divider()

    st.subheader("5) 지원 직무/지원 동기(선택)")
    target_role = st.text_input("지원 직무", value="교육공학 연구/기획")
    motivation = st.text_area(
        "지원 동기",
        value="저의 경험을 바탕으로 학습자 경험 중심의 원격교육 고도화와 AI 기반 교육혁신에 기여하고 싶습니다.",
        height=110,
    )

# -----------------------------
# Data object
# -----------------------------
profile = {
    "name": name,
    "role": role,
    "contact": {
        "email": email,
        "phone": phone,
        "location": location,
        "website": website,
    },
    "headline": headline,
    "intro": intro,
    "strengths": strengths,
    "skills": skills,
    "experiences": experiences,
    "target_role": target_role,
    "motivation": motivation,
}

# -----------------------------
# Preview renderers (값이 코드처럼 안 보이게)
# -----------------------------
def preview_card(data: dict):
    st.subheader("미리보기 (카드형)")

    with st.container(border=True):
        a, b = st.columns([1, 3], gap="medium")
        with a:
            if app_icon:
                st.image(app_icon, width=110)
            else:
                st.markdown("### 🧾")
        with b:
            st.markdown(f"## {data['name']}")
            st.markdown(f"**{data['role']}**")
            st.write(data["headline"])

            c1, c2, c3 = st.columns(3)
            c1.caption("📍 지역")
            c1.write(data["contact"]["location"] or "-")
            c2.caption("✉️ 이메일")
            c2.write(data["contact"]["email"] or "-")
            c3.caption("🔗 웹사이트")
            c3.write(data["contact"]["website"] or "-")

        st.markdown("### 자기소개")
        st.write(data["intro"])

        st.markdown("### 강점")
        st.markdown(data["strengths"])

        st.markdown("### 스킬")
        st.write(" / ".join(data["skills"]) if data["skills"] else "-")

    st.markdown("### 경험(경력/프로젝트)")
    if not data["experiences"]:
        st.info("입력된 경험이 없습니다.")
    else:
        for exp in data["experiences"]:
            with st.container(border=True):
                st.markdown(f"**{exp['title']}**")
                st.caption(f"{exp['org']} · {exp['from']} ~ {exp['to']}")
                st.write(exp["detail"])

    st.markdown("### 지원 직무 / 동기")
    with st.container(border=True):
        st.markdown(f"**지원 직무:** {data['target_role'] or '-'}")
        st.write(data["motivation"] or "-")


def preview_doc(data: dict):
    st.subheader("미리보기 (문서형)")

    st.markdown(f"# {data['name']}")
    st.markdown(f"**{data['role']}**")
    st.write(data["headline"])
    st.caption(
        f"📍 {data['contact']['location'] or '-'} | "
        f"✉️ {data['contact']['email'] or '-'} | "
        f"📞 {data['contact']['phone'] or '-'} | "
        f"🔗 {data['contact']['website'] or '-'}"
    )

    st.markdown("## 자기소개")
    st.write(data["intro"])

    st.markdown("## 강점")
    st.markdown(data["strengths"])

    st.markdown("## 스킬")
    st.write(", ".join(data["skills"]) if data["skills"] else "-")

    st.markdown("## 경험(경력/프로젝트)")
    if not data["experiences"]:
        st.write("-")
    else:
        for exp in data["experiences"]:
            st.markdown(f"### {exp['title']}")
            st.caption(f"{exp['org']} · {exp['from']} ~ {exp['to']}")
            st.write(exp["detail"])

    st.markdown("## 지원 직무 / 지원 동기")
    st.markdown(f"**지원 직무:** {data['target_role'] or '-'}")
    st.write(data["motivation"] or "-")


# -----------------------------
# Right panel
# -----------------------------
with right:
    if preview_mode == "카드형":
        preview_card(profile)
    else:
        preview_doc(profile)

    st.divider()

    st.subheader("내보내기")
    st.download_button(
        "📥 profile.json 다운로드",
        data=json.dumps(profile, ensure_ascii=False, indent=2).encode("utf-8"),
        file_name="profile.json",
        mime="application/json",
    )

    if show_json:
        st.subheader("입력값(JSON)")
        st.json(profile)

def make_docx(profile: dict) -> bytes:
    doc = Document()

    # 기본 폰트(선택): Word에서 한글 표시 안정성을 위해
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    doc.add_heading("자기소개서", level=1)

    # 기본 정보
    doc.add_paragraph(f"이름: {profile['name']}")
    doc.add_paragraph(f"직함/역할: {profile['role']}")
    contact = profile.get("contact", {})
    doc.add_paragraph(
        f"연락처: 이메일 {contact.get('email','-')} / "
        f"전화 {contact.get('phone','-')} / "
        f"지역 {contact.get('location','-')} / "
        f"웹사이트 {contact.get('website','-')}"
    )

    doc.add_paragraph("")  # 빈 줄
    doc.add_heading("한 줄 핵심 메시지", level=2)
    doc.add_paragraph(profile.get("headline", ""))

    doc.add_heading("자기소개", level=2)
    doc.add_paragraph(profile.get("intro", ""))

    doc.add_heading("강점", level=2)
    strengths_text = profile.get("strengths", "").strip()
    if strengths_text:
        # "- "로 시작하는 줄을 불릿 리스트로 변환
        for line in strengths_text.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("-"):
                doc.add_paragraph(line.lstrip("-").strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("-")

    doc.add_heading("스킬", level=2)
    skills = profile.get("skills", [])
    doc.add_paragraph(", ".join(skills) if skills else "-")

    doc.add_heading("경험(경력/프로젝트)", level=2)
    exps = profile.get("experiences", [])
    if not exps:
        doc.add_paragraph("-")
    else:
        for exp in exps:
            doc.add_heading(exp.get("title", ""), level=3)
            org = exp.get("org", "-")
            period = f"{exp.get('from','-')} ~ {exp.get('to','-')}"
            doc.add_paragraph(f"{org} · {period}")
            detail = exp.get("detail", "")
            if detail:
                doc.add_paragraph(detail)

    doc.add_heading("지원 직무", level=2)
    doc.add_paragraph(profile.get("target_role", "-") or "-")

    doc.add_heading("지원 동기", level=2)
    doc.add_paragraph(profile.get("motivation", "-") or "-")

    # bytes로 저장
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

docx_bytes = make_docx(profile)

st.download_button(
    "📝 자기소개서(.docx) 다운로드",
    data=docx_bytes,
    file_name=f"자기소개서_{profile['name']}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
